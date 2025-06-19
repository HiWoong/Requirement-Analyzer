import tempfile
import pandas as pd
import json
import os
import requests
import docx

from uuid import uuid4
from datetime import datetime, timedelta
from io import BytesIO
from collections import Counter, defaultdict
from openai import AzureOpenAI
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
from dotenv import load_dotenv
from azure.search.documents import SearchClient
from azure.data.tables import TableServiceClient
from azure.core.credentials import AzureKeyCredential

load_dotenv()
AZURE_STORAGE_CONNECTION_STRING = os.environ.get("AZURE_STORAGE_CONNECTION_STRING")
AZURE_STORAGE_ACCOUNT = os.environ.get("AZURE_STORAGE_ACCOUNT")
AZURE_STORAGE_KEY = os.environ.get("AZURE_STORAGE_KEY")
AZURE_BLOB_CONTAINER_NAME = os.environ.get("AZURE_BLOB_CONTAINER_NAME")

AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.environ.get("AZURE_OPENAI_KEY")
AZURE_OPENAI_DEPLOYMENT_EMBED = os.environ.get("AZURE_OPENAI_DEPLOYMENT_EMBED")
AZURE_OPENAI_DEPLOYMENT_EMBED_KEY = os.environ.get("AZURE_OPENAI_DEPLOYMENT_EMBED_KEY")
AZURE_OPENAI_DEPLOYMENT_GPT = os.environ.get("AZURE_OPENAI_DEPLOYMENT_GPT")

AZURE_SEARCH_ENDPOINT = os.environ.get("AZURE_SEARCH_ENDPOINT")
AZURE_SEARCH_KEY = os.environ.get("AZURE_SEARCH_KEY")
AZURE_SEARCH_INDEX_NAME = os.environ.get("AZURE_SEARCH_INDEX_NAME")

AZURE_TABLE_NAME = os.environ.get("AZURE_TABLE_NAME")

chat_client = AzureOpenAI(
    api_version = "2024-12-01-preview",
    azure_endpoint = AZURE_OPENAI_ENDPOINT,
    api_key = AZURE_OPENAI_KEY
)

client = AzureOpenAI(
        api_version="2024-12-01-preview",
        azure_deployment=AZURE_OPENAI_DEPLOYMENT_GPT,
        azure_endpoint = AZURE_OPENAI_ENDPOINT,
        api_key = AZURE_OPENAI_KEY,
    )

search_client = SearchClient(
    endpoint   = AZURE_SEARCH_ENDPOINT,
    index_name = AZURE_SEARCH_INDEX_NAME,
    credential = AzureKeyCredential(AZURE_SEARCH_KEY)
)

blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
container_client = blob_service_client.get_container_client(AZURE_BLOB_CONTAINER_NAME)

table_service = TableServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
try:
    table_client = table_service.create_table_if_not_exists(table_name=AZURE_TABLE_NAME)
except Exception as e:
    table_client = table_service.get_table_client(table_name=AZURE_TABLE_NAME)

# 테이블 저장
def save_to_table(key: str, value: str):
    entity = {
        "PartitionKey": "RTM",       # PartitionKey는 문자열 그룹 기준
        "RowKey": key,               # 고유 Key
        "Value": value
    }
    table_client.upsert_entity(entity)
    print(f"[저장 완료] {key} → {value}")

# 테이블 조회
def get_value_by_key(key: str) -> str:
    try:
        entity = table_client.get_entity(partition_key="RTM", row_key=key)
        return entity["Value"]
    except Exception as e:
        return key

def extract_requirements(text: str):
    prompt = [
        {
            "role": "system",
            "content": """
                You are a requirements extraction assistant.  
                Your task is to analyze the user's input (such as meeting notes or documentation) and extract clearly defined software requirements.
                Respond only with a valid JSON object that contains a key named `"items"` with an array of extracted requirements.  
                Each requirement must follow the exact structure below. Make sure no two items have duplicate descriptions.
                Although your internal reasoning can be in English, your final response must always be written in Korean.

                Each item must contain:
                - 요구사항 ID (Requirement ID): A unique ID in the format "REQ-001", "REQ-002", etc.
                - 요구사항 명 (Title): A concise phrase summarizing the requirement
                - 리스크 요인 (Risk Factor): Any potential technical or policy-related risk
                - 분석 코멘트 (Analysis Comment): Any useful insight or consideration for implementation
                - 설명 (Description): A full sentence or paragraph describing the requirement

                Example output format:
                {
                "items": [
                    {
                    "요구사항 ID": "REQ-001",
                    "요구사항 명": "사용자 인증",
                    "리스크 요인": "인증 충돌 가능성",
                    "분석 코멘트": "모듈화 리팩터링 권장",
                    "설명": "로그인 UX 개선 (지문/패턴 선택 제공)"
                    }
                ]
                }
                """
        },
        {
            "role": "user",
            "content": text
        }
    ]

    response = client.chat.completions.create(
        model = AZURE_OPENAI_DEPLOYMENT_GPT,
        messages = prompt,
        max_tokens=2000,
        temperature=0.7,
        response_format={"type": "json_object"} 
    )
    # 잘렸는지 확인
    if response.choices[0].finish_reason == "length":
        raise RuntimeError("GPT 응답이 잘렸습니다. max_tokens를 늘리세요.")
    return json.loads(response.choices[0].message.content.strip())["items"]

def upsert_rtm_rows(rtm_rows: list[dict], source_file: str):
    """요구사항을 Search Index에 벡터와 함께 저장"""
    actions = []
    columns = ["요구사항 ID", "설명", "난이도", "공수", "연관 회의록", "리스크 요인", "분석 코멘트", "상태"]
    header = '\t'.join(["Sheet1"] + columns)

    for idx, row in enumerate(rtm_rows, start=1):
        rows_as_text = [
            '\t' + '\t'.join(str(row.get(col, "")) for col in columns)
        ]
        content = header + '\n' + '\n'.join(rows_as_text)
        
        actions.append({
            "chunk_id": str(uuid4()),  # 또는 고유 ID 생성 방식
            "parent_id": row.get("요구사항 ID", f"REQ-{idx:03d}"),
            "chunk": content,
            "title": source_file,
            "text_vector": get_embedding(content)  # 1536‑float 배열
        })
    print("========================")
    print(len(actions))
    print("========================")
    result = search_client.merge_or_upload_documents(documents=actions)
    print(result)

# Azure OpenAI 임베딩 생성 함수
def get_embedding(text: str) -> list:
    client = AzureOpenAI(
        azure_deployment=AZURE_OPENAI_DEPLOYMENT_EMBED,
        api_version="2024-12-01-preview",
        azure_endpoint = AZURE_OPENAI_ENDPOINT,
        api_key = AZURE_OPENAI_KEY
        )
    response = client.embeddings.create(
        model=AZURE_OPENAI_DEPLOYMENT_EMBED,
        input=text
    )
    embeddings = response.data[0].embedding
    return embeddings

# Azure Cognitive Search 벡터 검색 함수 (REST API 직접 호출)
def vector_search(query_embedding: list, top_k=5) -> list:
    url = (
            f"{AZURE_SEARCH_ENDPOINT}/indexes/{AZURE_SEARCH_INDEX_NAME}/docs/search?api-version=2023-07-01-Preview"
        )

    headers = {
        "Content-Type": "application/json",
        "api-key": AZURE_SEARCH_KEY,
    }

    body = {
        "vector": {               # ★ 벡터 검색 조건
            "value": query_embedding,
            "k": top_k,
            "fields": "text_vector"  # 인덱스에 만든 벡터 필드명
        },
        "search": "*",            # 필수 placeholder
        "select": "*", 
        "top": top_k
    }

    resp = requests.post(url, headers=headers, data=json.dumps(body))
    resp.raise_for_status()       # 오류 시 예외

    # 결과 문서 리스트 반환
    return resp.json().get("value", [])

def detect_conflict_gpt(text1: str, text2: str) -> str:
    prompt = [
        {
            "role": "system",
            "content": """
                You are an expert in software requirements analysis.
                Your task is to compare the following two software requirements and classify their relationship into one of the following three categories:
                1. 중복 (Duplicate): The two requirements describe almost the same functionality.
                2. 충돌 (Conflict): The requirements are related but contradictory or mutually exclusive.
                3. 양호 (Unrelated): The requirements are unrelated or independent in scope.

                Your explanation ("사유") and final answer must be written in Korean.
                Return your answer strictly in the following JSON format:
                {"판단": "중복" | "충돌" | "양호", "사유": "A brief explanation of your reasoning"}
                """
                        },
                        {
                            "role": "user",
                            "content": f"""
                요구사항1: {text1}
                요구사항2: {text2}
            """
        }
    ]

    response = client.chat.completions.create(
        model=AZURE_OPENAI_DEPLOYMENT_GPT,
        messages=prompt,
        temperature=0,
        max_tokens=300,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content.strip())
    return data["판단"]

def extract_field_from_chunk(chunk: str, field_name: str) -> str:
    try:
        lines = chunk.strip().split('\n')
        if len(lines) < 2:
            return ""
        
        headers = lines[0].split('\t')
        values = lines[1].split('\t')
        
        if field_name in headers:
            index = headers.index(field_name)
            return values[index] if index < len(values) else ""
        else:
            return ""
    except Exception as e:
        return ""

def enrich_with_similarity(req_row: dict) -> dict:
    emb = get_embedding(req_row["설명"])
    hits = vector_search(emb, top_k=3)
    top_filename = get_value_by_key(hits[0].get("title", "") if hits else "")
    score = hits[0]["@search.score"] if hits else 0

    # GPT 2차 분석
    prompt = [
        {
            "role": "system",
            "content":
                f"""
                You are an expert software estimator.
                Based on the following requirement description, associated risk factor, analysis comment, and information from similar past documents, estimate and return the difficulty and expected development effort as a JSON object.
                Please follow the definitions below:
                - 난이도 (Difficulty): Choose one of ["쉬움", "보통", "어려움"]
                - 공수 (Effort): Estimated effort in MD (man-days), where 1 MD = 1 developer working for 1 day

                요구사항: {req_row['설명']}
                리스크 요인: {req_row['리스크 요인']}
                분석 코멘트: {req_row['분석 코멘트']}
                유사도score: {score}
                관련파일: {top_filename}
                """
        }
    ]

    response = client.chat.completions.create(
        model = AZURE_OPENAI_DEPLOYMENT_GPT,
        messages = prompt,
        max_tokens=300,
        temperature=0.7,
        response_format={"type": "json_object"}
    )
    data = json.loads(response.choices[0].message.content.strip())

    status = "양호"
    print(score)
    if score > 0.65:
        status = "중복"
    else:
        if 0.5 <= score <= 0.65 and hits:
            # 유사 문서 설명 추출
            chunk = hits[0].get("chunk", "")
            text2 = extract_field_from_chunk(chunk, field_name="설명")
            print(text2)
            conflict_result = detect_conflict_gpt(req_row["설명"], text2)
            if conflict_result == "충돌":
                status = "충돌"


    req_row.update({
        "난이도": data["난이도"],
        "공수": f"{data['공수']}MD",
        "상태": status,
        "연관 회의록": top_filename,
    })
    return req_row

def upload_blob(file_stream, filename: str) -> str:
    # 업로드
    if filename.endswith(".xlsx"):
        filename = f"rtm/{filename}"
    container_client.upload_blob(name=filename, data=file_stream, overwrite=True)
    sas_token = make_sas_token(filename)
    blob_url = f"https://{blob_service_client.account_name}.blob.core.windows.net/{AZURE_BLOB_CONTAINER_NAME}/{filename}?{sas_token}"
    return blob_url

# Azure Storage에 있는 특정 확장자의 파일 개수
def count_blob(ext: str) -> int:
    blob_list = container_client.list_blobs()
    count = sum(1 for blob in blob_list if blob.name.endswith(ext))
    return count

def find_most_relevant(final_rows: list[dict]) -> str:
    note_list = []
    for item in final_rows:
        note = item.get("연관 회의록")
        if note:
            note_list.append(note)

    if not note_list:
        return ""
    
    counter = Counter(note_list)
    most_relevant_note = counter.most_common(1)[0][0]
    return most_relevant_note

def count_recent_blob(ext: str) -> list:
    today = datetime.now().date()
    # 최근 7일 날짜 리스트 (6일 전부터 오늘까지)
    dates = [(today - timedelta(days=i)) for i in range(6, -1, -1)]
    date_strs = [d.strftime("%Y-%m-%d") for d in dates]

    # 날짜별 개수 초기화
    date_counts = defaultdict(int)

    # 컨테이너 내 모든 blob 이름과 생성일(또는 수정일) 확인
    # blob 생성일 정보가 없으면 blob.last_modified(UTC datetime)을 사용 가능
    blobs = container_client.list_blobs()
    for blob in blobs:
        # blob.last_modified은 UTC datetime
        blob_date = blob.last_modified.date()
        blob_date_str = blob_date.strftime("%Y-%m-%d")

        if blob_date_str in date_strs and blob.name.endswith(ext):
            date_counts[blob_date_str] += 1

    # dates 순서대로 개수 리스트 생성
    y = [date_counts[d] for d in date_strs]
    return y

def make_sas_token(blob:str) -> str:
    sas_token = generate_blob_sas(
        account_name=AZURE_STORAGE_ACCOUNT,
        container_name=AZURE_BLOB_CONTAINER_NAME,
        blob_name=blob,
        account_key=AZURE_STORAGE_KEY,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.now() + timedelta(hours=1)   # 1시간 유효
    )
    return sas_token

def get_openai_response(messages):
    # Azure AI Query용 파라미터 추가(Additional parameters to apply RAG pattern using the AI Search index)
    rag_params = {
        "data_sources" : [
            {
                "type" : "azure_search",
                "parameters" : {
                    "endpoint" : AZURE_SEARCH_ENDPOINT,
                    "index_name" : AZURE_SEARCH_INDEX_NAME,
                    "authentication" : {
                        "type" : "api_key",
                        "key" : AZURE_SEARCH_KEY
                    },
                    # 벡터 기반 검색
                    "query_type" : "vector",
                    # 입력을 벡터화(벡터화 시켜야만 비교 가능)
                    "embedding_dependency" : {
                        "type" : "deployment_name",
                        "deployment_name" : AZURE_OPENAI_DEPLOYMENT_EMBED
                    }
                }
            }
        ]
    }

    # RAG 파라미터와 함께 전송(Submit the chat request with RAG parameters)
    response = client.chat.completions.create(
        model = AZURE_OPENAI_DEPLOYMENT_GPT,
        messages = messages,
        extra_body = rag_params
    )

    completion = response.choices[0].message.content

    # citation 정보가 있는 경우 치환
    context = getattr(response.choices[0].message, "context", None)
    if context and "citations" in context:
        citations = context["citations"]
        doc_map = {
            f"[doc{i+1}]": f"[{c.get('title', f'doc{i+1}')}]"
            for i, c in enumerate(citations)
        }
        for key, value in doc_map.items():
            completion = completion.replace(key, value)
    return completion

def extract_text_from_file(uploaded_file):
    file_type = uploaded_file.name.split('.')[-1].lower()

    if file_type == 'txt':
        # txt 파일은 그대로 읽음
        text = uploaded_file.getvalue().decode('utf-8')
    elif file_type == 'docx':
        # docx 텍스트 추출
        doc = docx.Document(uploaded_file)
        paragraphs = [para.text for para in doc.paragraphs]
        text = '\n'.join(paragraphs)

    # 텍스트를 다시 BytesIO로 변환 (process_document에 txt와 동일한 인자 형태로 넘기기 위함)
    return BytesIO(text.encode('utf-8'))

# 가장 최근에 분석된 RTM 파일 가져오기
def load_latest_rtm_excel():
    # rtm/ 경로 아래에 있는 모든 .xlsx 파일 가져오기
    blob_list = container_client.list_blobs(name_starts_with="rtm/")
    xlsx_blobs = [
        blob for blob in blob_list
        if blob.name.endswith(".xlsx")
    ]

    # 아무것도 없으면 rtm_sample 반환
    if not xlsx_blobs:
        rtm_path = "data/rtm_sample.xlsx"
        return pd.read_excel(rtm_path), "", 0, ""

    # 가장 최근에 업로드된 파일 선택 (last_modified 기준)
    latest_blob = max(xlsx_blobs, key=lambda b: b.last_modified)

    # blob 다운로드
    blob_data = container_client.download_blob(latest_blob.name).readall()

    # BytesIO로 pandas에서 읽을 수 있게 변환
    rtm_df = pd.read_excel(BytesIO(blob_data))

    sas_token = make_sas_token(latest_blob.name)
    blob_url = f"https://{blob_service_client.account_name}.blob.core.windows.net/{AZURE_BLOB_CONTAINER_NAME}/{latest_blob.name}?{sas_token}"
    
    return rtm_df, blob_url, get_value_by_key(latest_blob.name.split("/")[-1]), find_most_relevant(rtm_df.to_dict(orient="records"))

# --- 요구사항 분석 및 RTM 생성 ---
def process_document(file_bytes: BytesIO, filename: str, created_by: str = "anonymous") -> pd.DataFrame:
    # 텍스트 추출
    text = file_bytes.getvalue().decode("utf-8")

    # GPT 1차 분석(요구사항 추출)
    # json array로 반환
    first_rtm = extract_requirements(text)

    # json array를 순회하며 벡터화 및 유사도 채워넣기
    final_rows = [enrich_with_similarity(r) for r in first_rtm]

    # DataFrame으로 변환
    df = pd.DataFrame(final_rows)
    
    # 고유 RTM 파일명 생성
    rtm_filename = f"RTM_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}.xlsx"

    # 테이블에 매핑
    save_to_table(filename, filename)
    save_to_table(rtm_filename, filename)
    
    # 엑셀 파일로 저장
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        df.to_excel(tmp.name, index=False)
        with open(tmp.name, "rb") as f:
            blob_url = upload_blob(f, rtm_filename)
    
    # 인덱스 업데이트
    upsert_rtm_rows(final_rows, filename)

    # txt 파일 저장
    upload_blob(file_bytes, filename)

    # 결과 반환
    return df, blob_url, count_blob(".txt"), find_most_relevant(final_rows)